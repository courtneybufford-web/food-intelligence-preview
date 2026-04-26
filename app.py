import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import requests
import re
import zipfile
import base64
from io import BytesIO

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None
    ImageDraw = None
    ImageFont = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
except Exception:
    colors = None
    letter = None
    landscape = None
    getSampleStyleSheet = None
    SimpleDocTemplate = None
    Table = None
    TableStyle = None
    Paragraph = None
    Spacer = None
try:
    from docx import Document
except Exception:
    Document = None
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

st.set_page_config(page_title="Food Intelligence App", layout="wide")
st.title("Food Intelligence App")
st.caption("Predictive database search: Customer database + sample database + USDA + Open Food Facts")

BADGES = {
    "Customer": "CUSTOMER",
    "Sample": "SAMPLE",
    "USDA": "USDA",
    "Open Food Facts": "OFF",
}

COMMON_UNITS = ["g", "kg", "mg", "oz", "lb", "ml", "l", "tsp", "tbsp", "cup", "fl oz", "pint", "quart", "gallon", "each", "piece", "slice", "serving", "portion"]
DEFAULT_LABEL_SIZE = "2 x 4 in"
DEFAULT_LABEL_DPI = 300

UNIT_TO_GRAMS = {
    "g": 1.0,
    "kg": 1000.0,
    "mg": 0.001,
    "oz": 28.3495,
    "lb": 453.592,
    "ml": 1.0,
    "l": 1000.0,
    "tsp": 4.93,
    "tbsp": 14.79,
    "cup": 240.0,
    "fl oz": 29.57,
    "pint": 473.0,
    "quart": 946.0,
    "gallon": 3785.0,
    "each": 0.0,
    "piece": 0.0,
    "slice": 0.0,
    "serving": 0.0,
    "portion": 0.0,
}

DEFAULT_PRODUCTS = [
    {"name":"Chicken Roti-Bulk", "source":"Sample", "calories":37274, "protein":410, "fat":620, "carbs":2800, "salt":42, "allergens":"", "ingredients":"chicken thigh, spices, oil, garlic, onion", "serving_note":"37274 cal per recipe yield"},
    {"name":"Harissa Chicken-Bulk", "source":"Sample", "calories":4372, "protein":360, "fat":190, "carbs":48, "salt":18, "allergens":"", "ingredients":"chicken, harissa paste, oil, garlic, lemon", "serving_note":"4372 cal per recipe yield"},
    {"name":"Hot Honey Harissa Chicken-3oz portion", "source":"Sample", "calories":209, "protein":22, "fat":8, "carbs":10, "salt":0.9, "allergens":"", "ingredients":"chicken, honey, harissa paste, spices", "serving_note":"209 cal per 3 oz portion"},
    {"name":"Chicken Roti Marinade", "source":"Sample", "calories":35665, "protein":12, "fat":3100, "carbs":120, "salt":64, "allergens":"mustard", "ingredients":"oil, lemon juice, garlic, mustard, spices", "serving_note":"35665 cal per recipe yield"},
    {"name":"Chicken, broilers or fryers, thigh, meat only, cooked, roasted", "source":"Sample", "calories":251, "protein":25, "fat":16, "carbs":0, "salt":0.22, "allergens":"", "ingredients":"chicken thigh meat", "serving_note":"251 cal per cup, chopped or diced"},
    {"name":"Pacific Organic Low Sodium Chicken Broth", "source":"Sample", "calories":10, "protein":1, "fat":0, "carbs":1, "salt":0.24, "allergens":"", "ingredients":"chicken broth, chicken flavor, sea salt", "serving_note":"10 cal per 240 ml"},
    {"name":"Wheat Bun", "source":"Sample", "calories":140, "protein":5, "fat":2, "carbs":26, "salt":0.55, "allergens":"cereals containing gluten", "ingredients":"wheat flour, water, yeast, salt", "serving_note":"140 cal per bun"},
    {"name":"Cheddar Cheese", "source":"Sample", "calories":113, "protein":7, "fat":9, "carbs":1, "salt":0.18, "allergens":"milk", "ingredients":"milk, salt, cultures, enzymes", "serving_note":"113 cal per slice"},
]

if "products" not in st.session_state:
    st.session_state.products = DEFAULT_PRODUCTS.copy()
if "recipe_items" not in st.session_state:
    st.session_state.recipe_items = []
if "saved_recipes" not in st.session_state:
    st.session_state.saved_recipes = []


def safe_float(value):
    try:
        if value is None or value == "":
            return 0.0
        return float(value)
    except Exception:
        return 0.0


def normalize_product(product):
    """Ensure every search result has the same keys."""
    return {
        "name": str(product.get("name", "Unnamed item")),
        "source": str(product.get("source", "Unknown")),
        "calories": safe_float(product.get("calories", 0)),
        "protein": safe_float(product.get("protein", 0)),
        "fat": safe_float(product.get("fat", 0)),
        "carbs": safe_float(product.get("carbs", 0)),
        "salt": safe_float(product.get("salt", 0)),
        "allergens": str(product.get("allergens", "") or ""),
        "ingredients": str(product.get("ingredients", "") or ""),
        "serving_note": str(product.get("serving_note", "") or f"{safe_float(product.get('calories', 0))} cal per serving"),
    }


def detect_allergens(text):
    t = (text or "").lower()
    hits = []
    rules = {
        "milk": ["milk", "cheese", "butter", "cream", "whey", "casein"],
        "cereals containing gluten": ["wheat", "barley", "rye", "oats", "spelt", "gluten"],
        "soybeans": ["soy", "soya", "soybean"],
        "eggs": ["egg"],
        "peanuts": ["peanut"],
        "nuts": ["almond", "cashew", "walnut", "pecan", "hazelnut", "pistachio"],
        "sesame": ["sesame"],
        "mustard": ["mustard"],
        "fish": ["fish"],
        "crustaceans": ["shrimp", "prawn", "crab", "lobster"],
        "molluscs": ["clam", "oyster", "mussel", "scallop"],
        "celery": ["celery"],
        "lupin": ["lupin"],
        "sulphites": ["sulfite", "sulphite", "sulfur dioxide", "sulphur dioxide"],
    }
    for allergen, words in rules.items():
        if any(w in t for w in words):
            hits.append(allergen)
    return ", ".join(sorted(set(hits)))



def extract_ingredients_from_text(text):
    raw = text or ""
    lower = raw.lower()
    if "ingredients" not in lower:
        return ""
    start = lower.find("ingredients")
    part = raw[start:]
    part = re.sub(r"^\s*ingredients\s*[:\-]?\s*", "", part.strip(), flags=re.IGNORECASE)
    stops = ["nutrition", "allergen", "contains", "storage", "preparation", "directions", "serving", "specification"]
    low_part = part.lower()
    positions = [low_part.find(stop) for stop in stops if low_part.find(stop) > 0]
    if positions:
        part = part[:min(positions)]
    return " ".join(part.replace("\n", " ").split()).strip(" .:")


def find_nutrient_value(text, labels):
    lower = (text or "").lower()
    for label in labels:
        pattern = rf"{re.escape(label.lower())}\D*(\d+(?:\.\d+)?)"
        match = re.search(pattern, lower)
        if match:
            return safe_float(match.group(1))
    return 0.0


def parse_product_spec_text(text, filename="Uploaded Spec"):
    raw = text or ""
    lines = [line.strip() for line in raw.splitlines() if line.strip()]
    guessed_name = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").strip().title()
    for line in lines[:8]:
        low = line.lower()
        if any(token in low for token in ["product name", "item name", "description"]):
            if ":" in line:
                guessed_name = line.split(":", 1)[1].strip() or guessed_name
                break
        elif len(line) <= 80 and not any(token in low for token in ["ingredients", "nutrition", "allergen", "contains"]):
            guessed_name = line
            break

    ingredients = extract_ingredients_from_text(raw)
    allergens = detect_allergens(raw + " " + ingredients)
    sodium_mg = find_nutrient_value(raw, ["sodium"])
    salt_g = find_nutrient_value(raw, ["salt"])
    if salt_g == 0 and sodium_mg:
        salt_g = round(sodium_mg * 2.5 / 1000, 3)
    calories = find_nutrient_value(raw, ["calories", "kcal", "energy"])
    product = normalize_product({
        "name": guessed_name,
        "source": "Customer",
        "calories": calories,
        "protein": find_nutrient_value(raw, ["protein"]),
        "fat": find_nutrient_value(raw, ["total fat", "fat"]),
        "carbs": find_nutrient_value(raw, ["total carbohydrate", "carbohydrate", "carbs"]),
        "salt": salt_g,
        "allergens": allergens,
        "ingredients": ingredients,
        "serving_note": f"{calories:g} cal per extracted serving" if calories else "Extracted from uploaded specification",
    })
    product["filename"] = filename
    product["raw_text"] = raw
    return product


def extract_text_from_pdf_bytes(file_bytes):
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""


def extract_text_from_file_bytes(name, data):
    lower = name.lower()
    if lower.endswith(".txt"):
        try:
            return data.decode("utf-8")
        except Exception:
            return data.decode("latin-1", errors="ignore")
    if lower.endswith(".pdf"):
        return extract_text_from_pdf_bytes(data)
    return ""


def parse_batch_uploads(uploaded_files):
    parsed = []
    errors = []
    for uploaded in uploaded_files:
        name = uploaded.name
        data = uploaded.read()
        lower = name.lower()
        if lower.endswith(".zip"):
            try:
                with zipfile.ZipFile(BytesIO(data)) as zf:
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        fname = info.filename
                        fl = fname.lower()
                        if not (fl.endswith(".txt") or fl.endswith(".pdf")):
                            continue
                        inner = zf.read(info)
                        extracted = extract_text_from_file_bytes(fname, inner)
                        if extracted.strip():
                            parsed.append(parse_product_spec_text(extracted, fname))
                        else:
                            errors.append(f"No text extracted from {fname}")
            except Exception as e:
                errors.append(f"Could not read ZIP {name}: {e}")
        elif lower.endswith(".txt") or lower.endswith(".pdf"):
            extracted = extract_text_from_file_bytes(name, data)
            if extracted.strip():
                parsed.append(parse_product_spec_text(extracted, name))
            else:
                errors.append(f"No text extracted from {name}")
        else:
            errors.append(f"Unsupported file type: {name}")
    return parsed, errors



def extract_text_from_docx_bytes(file_bytes):
    if Document is None:
        return ""
    try:
        doc = Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))
        return "\n".join(paragraphs + table_text)
    except Exception:
        return ""


def extract_text_from_excel_bytes(file_bytes, filename="uploaded.xlsx"):
    try:
        sheets = pd.read_excel(BytesIO(file_bytes), sheet_name=None, header=None)
        chunks = []
        for sheet_name, df in sheets.items():
            chunks.append(f"Sheet: {sheet_name}")
            for _, row in df.fillna("").iterrows():
                values = [str(v).strip() for v in row.tolist() if str(v).strip()]
                if values:
                    chunks.append(" | ".join(values))
        return "\n".join(chunks)
    except Exception:
        return ""


def extract_text_from_csv_bytes(file_bytes):
    try:
        df = pd.read_csv(BytesIO(file_bytes), header=None).fillna("")
        rows = []
        for _, row in df.iterrows():
            values = [str(v).strip() for v in row.tolist() if str(v).strip()]
            if values:
                rows.append(" | ".join(values))
        return "\n".join(rows)
    except Exception:
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            return file_bytes.decode("latin-1", errors="ignore")


def extract_text_from_recipe_file_bytes(name, data):
    lower = name.lower()
    if lower.endswith(".txt"):
        try:
            return data.decode("utf-8")
        except Exception:
            return data.decode("latin-1", errors="ignore")
    if lower.endswith(".pdf"):
        return extract_text_from_pdf_bytes(data)
    if lower.endswith(".docx"):
        return extract_text_from_docx_bytes(data)
    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        return extract_text_from_excel_bytes(data, name)
    if lower.endswith(".csv"):
        return extract_text_from_csv_bytes(data)
    return ""


def guess_recipe_name(text, filename="Uploaded Recipe"):
    base_name = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").strip().title()
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    for line in lines[:12]:
        low = line.lower()
        if any(label in low for label in ["recipe name", "menu item", "dish name", "product name"]):
            if ":" in line:
                return line.split(":", 1)[1].strip() or base_name
            if "|" in line:
                return line.split("|", 1)[-1].strip() or base_name
        if len(line) <= 80 and not any(token in low for token in ["ingredients", "method", "directions", "nutrition", "allergen", "yield", "servings"]):
            return line
    return base_name


def extract_servings_from_text(text):
    lower = (text or "").lower()
    patterns = [
        r"servings?\D+(\d+(?:\.\d+)?)",
        r"serves\D+(\d+(?:\.\d+)?)",
        r"yield\D+(\d+(?:\.\d+)?)",
        r"portions?\D+(\d+(?:\.\d+)?)",
    ]
    for pattern in patterns:
        m = re.search(pattern, lower)
        if m:
            return max(1, int(float(m.group(1))))
    return 1


def parse_recipe_upload_text(text, filename="Uploaded Recipe"):
    raw = text or ""
    name = guess_recipe_name(raw, filename)
    servings = extract_servings_from_text(raw)
    ingredients = extract_ingredients_from_text(raw)
    if not ingredients:
        rows = []
        for line in raw.splitlines():
            low = line.lower().strip()
            if not low or any(x in low for x in ["method", "direction", "instruction", "nutrition", "calories", "servings", "yield"]):
                continue
            if any(unit in low for unit in [" g", " kg", " oz", " lb", " cup", " tbsp", " tsp", " ml", " l "]):
                rows.append(line.strip())
        ingredients = ", ".join(rows[:40])
    if not ingredients:
        ingredients = "Needs manual ingredient review"
    calories = find_nutrient_value(raw, ["calories", "kcal", "energy"])
    protein = find_nutrient_value(raw, ["protein"])
    fat = find_nutrient_value(raw, ["total fat", "fat"])
    carbs = find_nutrient_value(raw, ["total carbohydrate", "carbohydrate", "carbs"])
    sodium_mg = find_nutrient_value(raw, ["sodium"])
    salt = find_nutrient_value(raw, ["salt"])
    if salt == 0 and sodium_mg:
        salt = round(sodium_mg * 2.5 / 1000, 3)
    per = {"calories": calories, "protein": protein, "fat": fat, "carbs": carbs, "salt": salt}
    detected_allergens = detect_allergens(raw + " " + ingredients)
    allergen_text = "Contains: " + (detected_allergens if detected_allergens else "No declarable allergens detected")
    label = f"""{name}\n\nIngredients: {ingredients}\n\n{allergen_text}\n\nNutrition per serving:\nEnergy: {per['calories']} kcal\nFat: {per['fat']} g\nCarbohydrate: {per['carbs']} g\nProtein: {per['protein']} g\nSalt: {per['salt']} g\n"""
    return {"name": name, "servings": servings, "items": [], "label": label, "nutrition_per_serving": per, "uploaded_ingredients": ingredients, "allergens": detected_allergens, "filename": filename, "raw_text": raw}


def parse_recipe_batch_uploads(uploaded_files):
    parsed = []
    errors = []
    supported = (".txt", ".pdf", ".docx", ".xlsx", ".xls", ".csv")
    for uploaded in uploaded_files:
        name = uploaded.name
        data = uploaded.read()
        lower = name.lower()
        if lower.endswith(".zip"):
            try:
                with zipfile.ZipFile(BytesIO(data)) as zf:
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        fname = info.filename
                        if not fname.lower().endswith(supported):
                            continue
                        inner = zf.read(info)
                        extracted = extract_text_from_recipe_file_bytes(fname, inner)
                        if extracted.strip():
                            parsed.append(parse_recipe_upload_text(extracted, fname))
                        else:
                            errors.append(f"No text extracted from {fname}")
            except Exception as e:
                errors.append(f"Could not read ZIP {name}: {e}")
        elif lower.endswith(supported):
            extracted = extract_text_from_recipe_file_bytes(name, data)
            if extracted.strip():
                parsed.append(parse_recipe_upload_text(extracted, name))
            else:
                errors.append(f"No text extracted from {name}")
        else:
            errors.append(f"Unsupported recipe file type: {name}")
    return parsed, errors


def item_grams(item):
    amount = safe_float(item.get("amount", item.get("qty", 1.0)))
    unit = item.get("unit", "serving")
    waste_pct = min(max(safe_float(item.get("waste_pct", 0.0)), 0.0), 100.0)
    grams = amount * UNIT_TO_GRAMS.get(unit, 0.0)
    usable_grams = grams * (1 - waste_pct / 100.0)
    return round(usable_grams, 2)


def nutrition_factor(item):
    amount = safe_float(item.get("amount", item.get("qty", 1.0)))
    waste_pct = min(max(safe_float(item.get("waste_pct", 0.0)), 0.0), 100.0)
    usable_amount = amount * (1 - waste_pct / 100.0)
    grams = item_grams(item)
    serving_note = str(item.get("serving_note", "")).lower()
    source = str(item.get("source", "")).lower()
    if grams > 0 and ("per 100 g" in serving_note or source in ["usda", "open food facts"]):
        return grams / 100.0
    return usable_amount

def totals(items):
    t = {"calories":0.0,"protein":0.0,"fat":0.0,"carbs":0.0,"salt":0.0}
    allergens = set()
    ingredients = []
    for item in items:
        factor = nutrition_factor(item)
        for k in t:
            t[k] += float(item.get(k, 0)) * factor
        if item.get("allergens"):
            for a in str(item["allergens"]).split(","):
                if a.strip():
                    allergens.add(a.strip())
        if item.get("ingredients"):
            ingredients.append(str(item["ingredients"]))
    return t, sorted(allergens), ", ".join(ingredients)



def scale_total_nutrition(total, factor):
    return {k: round(safe_float(v) * factor, 3) for k, v in total.items()}


def calculate_label_nutrition(total, servings, total_weight_g, serving_option, serving_size_value=1.0, serving_size_unit="serving", custom_serving_weight_g=0.0):
    """Return nutrition values to show on the label and the label serving text.

    The FDA-style default is per serving. Other options are provided for review,
    formulation, and non-retail/internal labels.
    """
    servings = max(int(servings or 1), 1)
    total_weight_g = safe_float(total_weight_g)
    serving_size_value = safe_float(serving_size_value) or 1.0
    serving_size_unit = serving_size_unit or "serving"

    if serving_option == "Per 100 g":
        factor = 100.0 / total_weight_g if total_weight_g > 0 else 0.0
        return scale_total_nutrition(total, factor), "100 g", servings

    if serving_option == "Per full recipe / container":
        return scale_total_nutrition(total, 1.0), "1 container", 1

    if serving_option == "Custom serving weight (g)":
        custom_g = safe_float(custom_serving_weight_g)
        factor = custom_g / total_weight_g if total_weight_g > 0 and custom_g > 0 else 1.0 / servings
        label = f"{custom_g:g} g" if custom_g > 0 else f"{serving_size_value:g} {serving_size_unit}"
        display_servings = round(total_weight_g / custom_g, 1) if total_weight_g > 0 and custom_g > 0 else servings
        return scale_total_nutrition(total, factor), label, display_servings

    # FDA/recommended default: per serving/container serving count.
    return scale_total_nutrition(total, 1.0 / servings), f"{serving_size_value:g} {serving_size_unit}", servings


def prediction_score(product, query):
    if not query:
        return 0
    q = query.lower().strip()
    name = product.get("name", "").lower()
    ingredients = product.get("ingredients", "").lower()
    source = product.get("source", "").lower()
    allergens = product.get("allergens", "").lower()
    haystack = f"{name} {ingredients} {source} {allergens}"
    if name.startswith(q):
        return 100
    if q in name:
        return 80
    if any(word.startswith(q) for word in name.split()):
        return 70
    if q in ingredients:
        return 55
    if q in haystack:
        return 40
    words = [w for w in q.split() if w]
    if words and all(w in haystack for w in words):
        return 35
    return 0


def search_customer_and_sample(query):
    scored = []
    for p in st.session_state.products:
        p = normalize_product(p)
        if not query:
            scored.append((1, p))
        else:
            score = prediction_score(p, query)
            if score > 0:
                scored.append((score, p))
    scored.sort(key=lambda x: (x[0], x[1].get("calories", 0)), reverse=True)
    return [p for _, p in scored]


@st.cache_data(ttl=3600, show_spinner=False)
def search_open_food_facts(query, limit=75):
    if not query or len(query.strip()) < 3:
        return []
    try:
        url = "https://world.openfoodfacts.org/cgi/search.pl"
        params = {
            "search_terms": query,
            "search_simple": 1,
            "action": "process",
            "json": 1,
            "page_size": limit,
            "fields": "product_name,generic_name,brands,ingredients_text,nutriments,allergens_tags"
        }
        data = requests.get(url, params=params, timeout=8).json()
        results = []
        for item in data.get("products", []):
            name = item.get("product_name") or item.get("generic_name") or "Open Food Facts item"
            nutriments = item.get("nutriments", {}) or {}
            ingredients = item.get("ingredients_text", "") or ""
            calories = safe_float(nutriments.get("energy-kcal_100g"))
            fat = safe_float(nutriments.get("fat_100g"))
            carbs = safe_float(nutriments.get("carbohydrates_100g"))
            protein = safe_float(nutriments.get("proteins_100g"))
            salt = safe_float(nutriments.get("salt_100g"))
            allergens = detect_allergens(ingredients + " " + " ".join(item.get("allergens_tags", [])))
            results.append(normalize_product({
                "name": name,
                "source": "Open Food Facts",
                "calories": calories,
                "protein": protein,
                "fat": fat,
                "carbs": carbs,
                "salt": salt,
                "allergens": allergens,
                "ingredients": ingredients,
                "serving_note": f"{calories:g} cal per 100 g" if calories else "Nutrition per 100 g when available",
            }))
        return results
    except Exception:
        return []


@st.cache_data(ttl=3600, show_spinner=False)
def search_usda(query, api_key, limit=75):
    if not api_key or not query or len(query.strip()) < 3:
        return []
    try:
        url = "https://api.nal.usda.gov/fdc/v1/foods/search"
        params = {"query": query, "pageSize": limit, "api_key": api_key}
        data = requests.get(url, params=params, timeout=8).json()
        results = []
        for food in data.get("foods", []):
            n = {"calories": 0, "protein": 0, "fat": 0, "carbs": 0, "salt": 0}
            for nutrient in food.get("foodNutrients", []):
                name = (nutrient.get("nutrientName") or "").lower()
                unit = (nutrient.get("unitName") or "").lower()
                value = safe_float(nutrient.get("value"))
                if "energy" in name and ("kcal" in unit or unit == "kcal"):
                    n["calories"] = value
                elif name == "protein":
                    n["protein"] = value
                elif name in ["total lipid (fat)", "total fat"]:
                    n["fat"] = value
                elif "carbohydrate" in name:
                    n["carbs"] = value
                elif name == "sodium, na":
                    n["salt"] = round(value * 2.5 / 1000, 3)
            desc = food.get("description", "USDA food").title()
            results.append(normalize_product({
                "name": desc,
                "source": "USDA",
                "calories": n["calories"],
                "protein": n["protein"],
                "fat": n["fat"],
                "carbs": n["carbs"],
                "salt": n["salt"],
                "allergens": detect_allergens(desc),
                "ingredients": desc.lower(),
                "serving_note": f"{n['calories']:g} cal per 100 g" if n["calories"] else "USDA nutrition per listed measure",
            }))
        return results
    except Exception:
        return []


def combined_database_search(query):
    api_key = st.secrets.get("USDA_API_KEY", "")
    local = search_customer_and_sample(query)
    usda = search_usda(query, api_key, limit=75) if len(query.strip()) >= 3 else []
    off = search_open_food_facts(query, limit=75) if len(query.strip()) >= 3 else []

    # Keep local matches first, then USDA, then OFF. Remove exact duplicate names from same source.
    seen = set()
    merged = []
    for item in local + usda + off:
        key = (item.get("source", ""), item.get("name", "").lower())
        if key not in seen:
            seen.add(key)
            merged.append(item)
    return merged, {"Customer/Sample": len(local), "USDA": len(usda), "Open Food Facts": len(off)}


def product_table(products):
    return pd.DataFrame([normalize_product(p) for p in products])


def badge_html(source):
    label = BADGES.get(source, source.upper() if source else "SOURCE")
    return f"<span style='font-size:11px; padding:2px 7px; border:1px solid #d8d2a8; border-radius:8px; background:#fbf7df; color:#776b22; font-weight:700;'>{label}</span>"


def render_search_result(product, index, prefix="search"):
    product = normalize_product(product)
    with st.container(border=True):
        icon_col, text_col, action_col, preview_col = st.columns([0.4, 5.7, 1.5, 1.2])
        with icon_col:
            if product.get("source") == "Customer":
                st.markdown("👤")
            elif product.get("source") == "USDA":
                st.markdown("✅")
            elif product.get("source") == "Open Food Facts":
                st.markdown("👥")
            else:
                st.markdown("🍽️")
        with text_col:
            st.markdown(f"**{product['name']}** &nbsp; {badge_html(product.get('source', 'Customer'))}", unsafe_allow_html=True)
            st.caption(product.get("serving_note") or f"{product.get('calories', 0):g} cal per serving")
            if product.get("allergens"):
                st.warning(f"Allergens: {product.get('allergens')}")
        with action_col:
            if st.button("+ Add to Recipe", key=f"{prefix}_add_{index}_{product['source']}_{product['name']}"):
                item = dict(product)
                item["amount"] = 1.0
                item["unit"] = "serving"
                item["waste_pct"] = 0.0
                item["grams"] = item_grams(item)
                st.session_state.recipe_items.append(item)
                st.success(f"Added {product['name']}")
        with preview_col:
            with st.expander("Preview"):
                st.write("**Ingredients**")
                st.write(product.get("ingredients", "") or "Not available")
                st.write("**Allergens**")
                st.write(product.get("allergens", "") or "none detected")
                st.write("**Nutrition**")
                st.json({
                    "calories": product.get("calories", 0),
                    "protein_g": product.get("protein", 0),
                    "fat_g": product.get("fat", 0),
                    "carbs_g": product.get("carbs", 0),
                    "salt_g": product.get("salt", 0),
                })




EXPORT_COLUMNS = [
    ("Menu Item", "name"),
    ("Serving Weight (g)", "serving_weight_g"),
    ("Calories", "calories"),
    ("Total Fat (g)", "fat"),
    ("Saturated Fat (g)", "saturated_fat"),
    ("Trans Fat (g)", "trans_fat"),
    ("Cholesterol (mg)", "cholesterol"),
    ("Sodium (mg)", "sodium"),
    ("Total Carbohydrate (g)", "carbs"),
    ("Dietary Fiber (g)", "fiber"),
    ("Total Sugars (g)", "sugars"),
    ("Added Sugars (g)", "added_sugars"),
    ("Protein (g)", "protein"),
    ("Salt (g)", "salt"),
    ("Ingredients", "ingredients"),
    ("Allergens", "allergens"),
]

def recipe_export_record(recipe):
    items = recipe.get("items", [])
    total, allergens, ingredients = totals(items)
    servings = max(safe_float(recipe.get("servings", 1)), 1.0)
    per = {k: round(v / servings, 3) for k, v in total.items()}
    serving_weight = round(sum(item_grams(item) for item in items) / servings, 2)

    return {
        "name": recipe.get("name", ""),
        "serving_weight_g": serving_weight,
        "calories": per.get("calories", 0),
        "fat": per.get("fat", 0),
        "saturated_fat": per.get("saturated_fat", 0),
        "trans_fat": per.get("trans_fat", 0),
        "cholesterol": per.get("cholesterol", 0),
        "sodium": per.get("sodium", 0),
        "carbs": per.get("carbs", 0),
        "fiber": per.get("fiber", 0),
        "sugars": per.get("sugars", 0),
        "added_sugars": per.get("added_sugars", 0),
        "protein": per.get("protein", 0),
        "salt": per.get("salt", 0),
        "ingredients": ingredients,
        "allergens": ", ".join(allergens) if isinstance(allergens, list) else str(allergens),
    }

def build_nutrition_export_dataframe(recipes):
    rows = []
    for recipe in recipes:
        record = recipe_export_record(recipe)
        rows.append({label: record.get(key, "") for label, key in EXPORT_COLUMNS})
    return pd.DataFrame(rows)

def build_allergen_export_dataframe(recipes):
    rows = []
    for recipe in recipes:
        record = recipe_export_record(recipe)
        rows.append({"Menu Item": record["name"], "Allergens": record["allergens"]})
    return pd.DataFrame(rows)

def build_ingredient_export_dataframe(recipes):
    rows = []
    for recipe in recipes:
        record = recipe_export_record(recipe)
        rows.append({"Menu Item": record["name"], "Ingredients": record["ingredients"]})
    return pd.DataFrame(rows)

def build_diet_export_dataframe(recipes):
    rows = []
    for recipe in recipes:
        allergens = recipe_export_record(recipe)["allergens"].lower()
        ingredients = recipe_export_record(recipe)["ingredients"].lower()
        rows.append({
            "Menu Item": recipe.get("name", ""),
            "GF": "" if any(x in allergens + " " + ingredients for x in ["wheat", "gluten", "barley", "rye"]) else "x",
            "Vegan": "" if any(x in allergens + " " + ingredients for x in ["milk", "egg", "fish", "shrimp", "chicken", "beef", "pork", "honey"]) else "x",
            "Halal": "Review",
        })
    return pd.DataFrame(rows)

def create_excel_export(recipes):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        build_nutrition_export_dataframe(recipes).to_excel(writer, sheet_name="Nutritionals", index=False)
        build_allergen_export_dataframe(recipes).to_excel(writer, sheet_name="Allergens", index=False)
        build_ingredient_export_dataframe(recipes).to_excel(writer, sheet_name="Ingredients", index=False)
        build_diet_export_dataframe(recipes).to_excel(writer, sheet_name="GF VG HALAL", index=False)

        workbook = writer.book
        for sheet_name in writer.sheets:
            ws = writer.sheets[sheet_name]
            for col in ws.columns:
                header = str(col[0].value or "")
                if header in ["Ingredients", "Allergens"]:
                    width = 75
                elif header == "Menu Item":
                    width = 30
                else:
                    width = 18
                ws.column_dimensions[col[0].column_letter].width = width
            for cell in ws[1]:
                cell.font = cell.font.copy(bold=True)
            ws.freeze_panes = "A2"
    output.seek(0)
    return output.getvalue()

def create_pdf_export(recipes):
    if SimpleDocTemplate is None:
        return None

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(letter),
        rightMargin=18,
        leftMargin=18,
        topMargin=18,
        bottomMargin=18,
    )
    styles = getSampleStyleSheet()
    story = [Paragraph("Full Nutrition Export", styles["Title"]), Spacer(1, 12)]

    df = build_nutrition_export_dataframe(recipes)
    short_cols = [
        "Menu Item", "Serving Weight (g)", "Calories", "Total Fat (g)",
        "Saturated Fat (g)", "Cholesterol (mg)", "Sodium (mg)",
        "Total Carbohydrate (g)", "Dietary Fiber (g)", "Total Sugars (g)",
        "Protein (g)", "Salt (g)", "Allergens"
    ]
    table_data = [short_cols]
    for _, row in df.iterrows():
        table_data.append([str(row.get(c, ""))[:80] for c in short_cols])

    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    story.append(Spacer(1, 14))
    story.append(Paragraph("Ingredient Statements", styles["Heading2"]))

    for recipe in recipes:
        record = recipe_export_record(recipe)
        story.append(Paragraph(f"<b>{record['name']}</b>", styles["Heading3"]))
        story.append(Paragraph(f"Ingredients: {record['ingredients']}", styles["BodyText"]))
        story.append(Paragraph(f"Allergens: {record['allergens']}", styles["BodyText"]))
        story.append(Spacer(1, 8))

    doc.build(story)
    output.seek(0)
    return output.getvalue()



FDA_DV = {
    "total_fat_g": 78,
    "saturated_fat_g": 20,
    "cholesterol_mg": 300,
    "sodium_mg": 2300,
    "total_carbs_g": 275,
    "dietary_fiber_g": 28,
    "added_sugars_g": 50,
    "protein_g": 50,
    "vitamin_d_mcg": 20,
    "calcium_mg": 1300,
    "iron_mg": 18,
    "potassium_mg": 4700,
}


def _nf_value(per_serving, *keys):
    """Read nutrient values from either older app keys or expanded FDA keys."""
    for key in keys:
        if key in per_serving and per_serving.get(key) not in [None, ""]:
            return safe_float(per_serving.get(key))
    return 0.0


def _round_nearest(value, increment):
    value = safe_float(value)
    return round(round(value / increment) * increment, 3) if increment else value


def _fmt_number(value):
    value = safe_float(value)
    if abs(value - int(value)) < 0.0001:
        return str(int(value))
    return str(round(value, 1)).rstrip("0").rstrip(".")


def fda_round_calories(calories):
    calories = safe_float(calories)
    if calories < 5:
        return 0
    if calories <= 50:
        return int(_round_nearest(calories, 5))
    return int(_round_nearest(calories, 10))


def fda_round_fat_grams(value):
    value = safe_float(value)
    if value < 0.5:
        return "0g"
    if value < 5:
        return f"{_fmt_number(_round_nearest(value, 0.5))}g"
    return f"{int(_round_nearest(value, 1))}g"


def fda_round_carb_grams(value, allow_less_than=True):
    value = safe_float(value)
    if value < 0.5:
        return "0g"
    if value < 1 and allow_less_than:
        return "<1g"
    return f"{int(_round_nearest(value, 1))}g"


def fda_round_cholesterol(value):
    value = safe_float(value)
    if value < 2:
        return "0mg"
    if value <= 5:
        return "<5mg"
    return f"{int(_round_nearest(value, 5))}mg"


def fda_round_sodium(value):
    value = safe_float(value)
    if value < 5:
        return "0mg"
    if value <= 140:
        return f"{int(_round_nearest(value, 5))}mg"
    return f"{int(_round_nearest(value, 10))}mg"


def fda_round_mineral(value, unit):
    value = safe_float(value)
    if value == 0:
        return f"0{unit}"
    if unit == "mcg":
        return f"{_fmt_number(_round_nearest(value, 0.1))}{unit}"
    return f"{_fmt_number(_round_nearest(value, 1))}{unit}"


def pdv(value, dv):
    value = safe_float(value)
    if dv <= 0:
        return ""
    return f"{int(round(value / dv * 100))}%"


def build_fda_nutrients(per_serving):
    sodium_mg = _nf_value(per_serving, "sodium_mg")
    if sodium_mg == 0:
        salt_g = _nf_value(per_serving, "salt", "salt_g")
        sodium_mg = salt_g * 1000 / 2.5 if salt_g else 0

    return {
        "calories": _nf_value(per_serving, "calories"),
        "total_fat_g": _nf_value(per_serving, "total_fat_g", "fat"),
        "saturated_fat_g": _nf_value(per_serving, "saturated_fat_g", "saturated_fat"),
        "trans_fat_g": _nf_value(per_serving, "trans_fat_g", "trans_fat"),
        "cholesterol_mg": _nf_value(per_serving, "cholesterol_mg", "cholesterol"),
        "sodium_mg": sodium_mg,
        "total_carbs_g": _nf_value(per_serving, "total_carbs_g", "carbs"),
        "dietary_fiber_g": _nf_value(per_serving, "dietary_fiber_g", "fiber"),
        "total_sugars_g": _nf_value(per_serving, "total_sugars_g", "sugars"),
        "added_sugars_g": _nf_value(per_serving, "added_sugars_g", "added_sugars"),
        "protein_g": _nf_value(per_serving, "protein_g", "protein"),
        "vitamin_d_mcg": _nf_value(per_serving, "vitamin_d_mcg", "vitamin_d"),
        "calcium_mg": _nf_value(per_serving, "calcium_mg", "calcium"),
        "iron_mg": _nf_value(per_serving, "iron_mg", "iron"),
        "potassium_mg": _nf_value(per_serving, "potassium_mg", "potassium"),
    }


def nutrition_facts_panel_text(recipe_name, per_serving, servings=1, serving_weight_g=0, serving_size_label=None):
    n = build_fda_nutrients(per_serving)
    serving_line = serving_size_label or (f"{serving_weight_g:g} g" if serving_weight_g else "1 serving")
    return f"""Nutrition Facts
{recipe_name}
{servings} servings per container
Serving size {serving_line}

Amount per serving
Calories {fda_round_calories(n['calories'])}

% Daily Value*
Total Fat {fda_round_fat_grams(n['total_fat_g'])} {pdv(n['total_fat_g'], FDA_DV['total_fat_g'])}
  Saturated Fat {fda_round_fat_grams(n['saturated_fat_g'])} {pdv(n['saturated_fat_g'], FDA_DV['saturated_fat_g'])}
  Trans Fat {fda_round_fat_grams(n['trans_fat_g'])}
Cholesterol {fda_round_cholesterol(n['cholesterol_mg'])} {pdv(n['cholesterol_mg'], FDA_DV['cholesterol_mg'])}
Sodium {fda_round_sodium(n['sodium_mg'])} {pdv(n['sodium_mg'], FDA_DV['sodium_mg'])}
Total Carbohydrate {fda_round_carb_grams(n['total_carbs_g'])} {pdv(n['total_carbs_g'], FDA_DV['total_carbs_g'])}
  Dietary Fiber {fda_round_carb_grams(n['dietary_fiber_g'])} {pdv(n['dietary_fiber_g'], FDA_DV['dietary_fiber_g'])}
  Total Sugars {fda_round_carb_grams(n['total_sugars_g'])}
    Includes {fda_round_carb_grams(n['added_sugars_g'])} Added Sugars {pdv(n['added_sugars_g'], FDA_DV['added_sugars_g'])}
Protein {fda_round_carb_grams(n['protein_g'], allow_less_than=False)}

Vitamin D {fda_round_mineral(n['vitamin_d_mcg'], 'mcg')} {pdv(n['vitamin_d_mcg'], FDA_DV['vitamin_d_mcg'])}
Calcium {fda_round_mineral(n['calcium_mg'], 'mg')} {pdv(n['calcium_mg'], FDA_DV['calcium_mg'])}
Iron {fda_round_mineral(n['iron_mg'], 'mg')} {pdv(n['iron_mg'], FDA_DV['iron_mg'])}
Potassium {fda_round_mineral(n['potassium_mg'], 'mg')} {pdv(n['potassium_mg'], FDA_DV['potassium_mg'])}

* The % Daily Value tells you how much a nutrient in a serving of food contributes to a daily diet. 2,000 calories a day is used for general nutrition advice.
Compliance note: FDA-style draft only. Verify serving size, RACC/category, lab/database values, rounding, ingredient statement, allergen declaration, and claims before commercial label use.
"""


def render_nutrition_facts_panel(recipe_name, per_serving, servings=1, serving_weight_g=0, serving_size_label=None):
    n = build_fda_nutrients(per_serving)
    serving_line = serving_size_label or (f"{serving_weight_g:g} g" if serving_weight_g else "1 serving")

    def row(label, amount, dv="", bold=False, indent=0, italic_first=False):
        fw = "800" if bold else "400"
        style = f"border-bottom:1px solid #111;padding:3px 0 3px {indent}px;display:flex;justify-content:space-between;font-size:13px;"
        if italic_first and " " in label:
            first, rest = label.split(" ", 1)
            label_html = f"<i>{first}</i> {rest}"
        else:
            label_html = label
        return f"<div style='{style}'><span style='font-weight:{fw};'>{label_html} {amount}</span><span style='font-weight:800;'>{dv}</span></div>"

    html = f"""
    <div style="max-width:390px;border:3.5px solid #111;padding:9px;background:white;color:#111;font-family:Arial, Helvetica, sans-serif;">
        <div style="font-weight:900;font-size:38px;line-height:36px;border-bottom:10px solid #111;letter-spacing:-1px;">Nutrition Facts</div>
        <div style="font-size:13px;margin-top:4px;"><b>{recipe_name}</b></div>
        <div style="font-size:13px;display:flex;justify-content:space-between;"><span>{servings} servings per container</span></div>
        <div style="font-size:15px;border-bottom:7px solid #111;display:flex;justify-content:space-between;"><b>Serving size</b><b>{serving_line}</b></div>
        <div style="font-size:12px;font-weight:800;margin-top:3px;">Amount per serving</div>
        <div style="display:flex;justify-content:space-between;align-items:flex-end;border-bottom:10px solid #111;">
            <div style="font-weight:900;font-size:29px;line-height:32px;">Calories</div>
            <div style="font-weight:900;font-size:38px;line-height:40px;">{fda_round_calories(n['calories'])}</div>
        </div>
        <div style="font-size:12px;text-align:right;border-bottom:1px solid #111;font-weight:800;">% Daily Value*</div>
        {row('Total Fat', fda_round_fat_grams(n['total_fat_g']), pdv(n['total_fat_g'], FDA_DV['total_fat_g']), True)}
        {row('Saturated Fat', fda_round_fat_grams(n['saturated_fat_g']), pdv(n['saturated_fat_g'], FDA_DV['saturated_fat_g']), False, 15)}
        {row('Trans Fat', fda_round_fat_grams(n['trans_fat_g']), '', False, 15, True)}
        {row('Cholesterol', fda_round_cholesterol(n['cholesterol_mg']), pdv(n['cholesterol_mg'], FDA_DV['cholesterol_mg']), True)}
        {row('Sodium', fda_round_sodium(n['sodium_mg']), pdv(n['sodium_mg'], FDA_DV['sodium_mg']), True)}
        {row('Total Carbohydrate', fda_round_carb_grams(n['total_carbs_g']), pdv(n['total_carbs_g'], FDA_DV['total_carbs_g']), True)}
        {row('Dietary Fiber', fda_round_carb_grams(n['dietary_fiber_g']), pdv(n['dietary_fiber_g'], FDA_DV['dietary_fiber_g']), False, 15)}
        {row('Total Sugars', fda_round_carb_grams(n['total_sugars_g']), '', False, 15)}
        {row('Includes', fda_round_carb_grams(n['added_sugars_g']) + ' Added Sugars', pdv(n['added_sugars_g'], FDA_DV['added_sugars_g']), False, 30)}
        <div style="border-bottom:8px solid #111;padding:3px 0;font-size:13px;"><b>Protein</b> {fda_round_carb_grams(n['protein_g'], allow_less_than=False)}</div>
        {row('Vitamin D', fda_round_mineral(n['vitamin_d_mcg'], 'mcg'), pdv(n['vitamin_d_mcg'], FDA_DV['vitamin_d_mcg']))}
        {row('Calcium', fda_round_mineral(n['calcium_mg'], 'mg'), pdv(n['calcium_mg'], FDA_DV['calcium_mg']))}
        {row('Iron', fda_round_mineral(n['iron_mg'], 'mg'), pdv(n['iron_mg'], FDA_DV['iron_mg']))}
        {row('Potassium', fda_round_mineral(n['potassium_mg'], 'mg'), pdv(n['potassium_mg'], FDA_DV['potassium_mg']))}
        <div style="border-top:5px solid #111;font-size:10px;line-height:12px;margin-top:4px;padding-top:4px;">* The % Daily Value tells you how much a nutrient in a serving of food contributes to a daily diet. 2,000 calories a day is used for general nutrition advice.</div>
        <div style="font-size:9px;margin-top:5px;color:#555;">Draft label. Verify serving size, lab/database values, rounding, ingredients, allergens, and claims before commercial use.</div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)


def copy_button(label, text, key="copy_btn"):
    import json
    safe_text = json.dumps(text)
    safe_label = label.replace("'", "&apos;")
    components.html(
        f"""
        <button style="background:#111;color:white;border:none;padding:8px 12px;border-radius:6px;font-weight:700;cursor:pointer;"
            onclick='navigator.clipboard.writeText({safe_text}); this.innerText="Copied!"; setTimeout(() => this.innerText="{safe_label}", 1400);'>
            {safe_label}
        </button>
        """,
        height=45,
    )



def create_nutrition_facts_pdf(recipe_name, panel_text):
    if SimpleDocTemplate is None:
        return None
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story = [Paragraph("Nutrition Facts Panel - FDA Style Draft", styles["Title"]), Spacer(1, 12)]

    table_data = [[line] if line.strip() else [""] for line in panel_text.splitlines()]
    table = Table(table_data, colWidths=[360])
    table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 2, colors.black),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("FONTNAME", (0, 0), (0, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (0, 0), 18),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    story.append(Spacer(1, 10))
    story.append(Paragraph("Draft compliance aid only. Verify against FDA regulations and product-specific lab/database data before commercial use.", styles["BodyText"]))
    doc.build(story)
    output.seek(0)
    return output.getvalue()




def _load_label_font(size, bold=False, condensed=False):
    if ImageFont is None:
        return None
    candidates = []
    if condensed:
        candidates += [
            "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
            "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        ]
    candidates += [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for fp in candidates:
        try:
            return ImageFont.truetype(fp, size)
        except Exception:
            pass
    return ImageFont.load_default()


def _draw_text_fit(draw, xy, text, font, fill="black", max_width=None, min_size=8, bold=False, condensed=True):
    if max_width is None or ImageFont is None:
        draw.text(xy, text, fill=fill, font=font)
        return font
    size = getattr(font, "size", min_size)
    active_font = font
    while size > min_size and draw.textlength(text, font=active_font) > max_width:
        size -= 1
        active_font = _load_label_font(size, bold=bold, condensed=condensed)
    draw.text(xy, text, fill=fill, font=active_font)
    return active_font


def _parse_panel_text(panel_text):
    lines = [line.strip() for line in (panel_text or "").splitlines() if line.strip()]
    parsed = {"recipe": "", "servings": "", "serving_size": "1 serving", "calories": "0", "rows": [], "vitamins": [], "footnote": ""}
    if len(lines) > 1 and lines[1].lower() != "nutrition facts":
        parsed["recipe"] = lines[1]
    for line in lines:
        low = line.lower()
        if "servings per container" in low:
            parsed["servings"] = line
        elif low.startswith("serving size"):
            parsed["serving_size"] = line.replace("Serving size", "").strip() or "1 serving"
        elif low.startswith("calories"):
            parts = line.split()
            parsed["calories"] = parts[-1] if parts else "0"
        elif low.startswith(("total fat", "saturated fat", "trans fat", "cholesterol", "sodium", "total carbohydrate", "dietary fiber", "total sugars", "includes", "protein")):
            m = re.search(r"\s(\d+%|<\s*\d+%)$", line)
            dv = m.group(1) if m else ""
            left = line[:m.start()].strip() if m else line
            parsed["rows"].append((left, dv))
        elif low.startswith(("vitamin d", "calcium", "iron", "potassium")):
            m = re.search(r"\s(\d+%|<\s*\d+%)$", line)
            dv = m.group(1) if m else ""
            left = line[:m.start()].strip() if m else line
            parsed["vitamins"].append((left, dv))
        elif low.startswith("* the % daily"):
            parsed["footnote"] = line
    if not parsed["servings"]:
        parsed["servings"] = "1 serving per container"
    if not parsed["footnote"]:
        parsed["footnote"] = "* The % Daily Value tells you how much a nutrient in a serving of food contributes to a daily diet."
    return parsed


def create_nutrition_facts_png(panel_text, label_size="2 x 4 in", dpi=300):
    # Create a high-resolution FDA-style Nutrition Facts PNG.
    if Image is None:
        return None
    size_map = {"2 x 4 in": (2, 4), "3 x 5 in": (3, 5), "4 x 6 in": (4, 6), "5 x 7 in": (5, 7)}
    width_in, height_in = size_map.get(label_size, (2, 4))
    width, height = int(width_in * dpi), int(height_in * dpi)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    margin = max(int(0.055 * dpi), 10)
    pad = max(int(0.035 * dpi), 7)
    left = margin + pad
    right = width - margin - pad
    y = margin + pad
    scale = max(width_in / 2.0, 1.0)
    title_font = _load_label_font(max(int(34 * scale), 28), True, True)
    serving_font = _load_label_font(max(int(8.8 * scale), 8), False, True)
    serving_bold = _load_label_font(max(int(9.3 * scale), 8), True, True)
    amount_font = _load_label_font(max(int(8 * scale), 7), True, True)
    calories_font = _load_label_font(max(int(22 * scale), 18), True, True)
    calories_num_font = _load_label_font(max(int(31 * scale), 24), True, True)
    row_font = _load_label_font(max(int(8.8 * scale), 8), False, True)
    row_bold = _load_label_font(max(int(8.8 * scale), 8), True, True)
    foot_font = _load_label_font(max(int(6.1 * scale), 6), False, True)
    parsed = _parse_panel_text(panel_text)
    border_w = max(int(0.012 * dpi), 3)
    thin_w = max(int(0.0035 * dpi), 1)
    med_w = max(int(0.009 * dpi), 2)
    heavy_w = max(int(0.030 * dpi), 7)
    draw.rectangle([margin, margin, width - margin, height - margin], outline="black", width=border_w)
    _draw_text_fit(draw, (left, y), "Nutrition Facts", title_font, max_width=right-left, min_size=20, bold=True)
    bbox = draw.textbbox((left, y), "Nutrition Facts", font=title_font)
    y = bbox[3] + max(int(0.010 * dpi), 3)
    draw.line([left, y, right, y], fill="black", width=heavy_w)
    y += max(int(0.019 * dpi), 5)
    if parsed["recipe"]:
        _draw_text_fit(draw, (left, y), parsed["recipe"], serving_bold, max_width=right-left, min_size=7, bold=True)
        y += max(int(0.055 * dpi), 14)
    draw.text((left, y), parsed["servings"], fill="black", font=serving_font)
    y += max(int(0.050 * dpi), 13)
    draw.text((left, y), "Serving size", fill="black", font=serving_bold)
    ss_width = draw.textlength(parsed["serving_size"], font=serving_bold)
    draw.text((right - ss_width, y), parsed["serving_size"], fill="black", font=serving_bold)
    y += max(int(0.060 * dpi), 16)
    draw.line([left, y, right, y], fill="black", width=heavy_w)
    y += max(int(0.018 * dpi), 5)
    draw.text((left, y), "Amount per serving", fill="black", font=amount_font)
    y += max(int(0.046 * dpi), 12)
    draw.text((left, y), "Calories", fill="black", font=calories_font)
    cal_text = str(parsed["calories"])
    cal_w = draw.textlength(cal_text, font=calories_num_font)
    draw.text((right - cal_w, y - max(int(0.020 * dpi), 5)), cal_text, fill="black", font=calories_num_font)
    y += max(int(0.095 * dpi), 26)
    draw.line([left, y, right, y], fill="black", width=heavy_w)
    y += max(int(0.010 * dpi), 3)
    dv = "% Daily Value*"
    dv_w = draw.textlength(dv, font=amount_font)
    draw.text((right - dv_w, y), dv, fill="black", font=amount_font)
    y += max(int(0.043 * dpi), 11)
    draw.line([left, y, right, y], fill="black", width=thin_w)

    def draw_row(label_amount, dv_text="", indent=0, bold=False):
        nonlocal y
        y += max(int(0.010 * dpi), 2)
        font = row_bold if bold else row_font
        x = left + indent
        if dv_text:
            dv_width = draw.textlength(dv_text, font=row_bold)
            max_left_width = right - x - dv_width - max(int(0.035 * dpi), 7)
        else:
            dv_width = 0
            max_left_width = right - x
        _draw_text_fit(draw, (x, y), label_amount, font, max_width=max_left_width, min_size=6, bold=bold)
        if dv_text:
            draw.text((right - dv_width, y), dv_text, fill="black", font=row_bold)
        y += max(int(0.044 * dpi), 11)
        draw.line([left + indent, y, right, y], fill="black", width=thin_w)

    for left_text, dv_text in parsed["rows"]:
        low = left_text.lower()
        indent = 0
        bold = low.startswith(("total fat", "cholesterol", "sodium", "total carbohydrate", "protein"))
        if low.startswith(("saturated", "trans", "dietary", "total sugars")):
            indent = max(int(0.075 * dpi), 14)
        if low.startswith("includes"):
            indent = max(int(0.145 * dpi), 26)
        draw_row(left_text, dv_text, indent=indent, bold=bold)
        if low.startswith("protein"):
            draw.line([left, y, right, y], fill="black", width=med_w)
            y += max(int(0.006 * dpi), 2)
    if parsed["vitamins"]:
        draw.line([left, y, right, y], fill="black", width=med_w)
        for left_text, dv_text in parsed["vitamins"]:
            draw_row(left_text, dv_text, indent=0, bold=False)
    foot_y = min(y + max(int(0.025 * dpi), 6), height - margin - pad - max(int(0.30 * dpi), 55))
    draw.line([left, foot_y, right, foot_y], fill="black", width=med_w)
    foot_y += max(int(0.012 * dpi), 3)
    words = parsed["footnote"].split()
    line = ""
    for word in words:
        test = (line + " " + word).strip()
        if draw.textlength(test, font=foot_font) <= right - left:
            line = test
        else:
            draw.text((left, foot_y), line, fill="black", font=foot_font)
            foot_y += max(int(0.034 * dpi), 8)
            line = word
    if line:
        draw.text((left, foot_y), line, fill="black", font=foot_font)
    output = BytesIO()
    img.save(output, format="PNG", dpi=(dpi, dpi))
    return output.getvalue()


def image_clipboard_tools(png_bytes, key="copy_image"):
    if not png_bytes:
        return
    b64 = base64.b64encode(png_bytes).decode("utf-8")
    components.html(f'''
        <div style="display:flex; gap:8px; align-items:center; flex-wrap:wrap; font-family:Arial, sans-serif;">
            <button id="copy_{key}" style="background:#0f172a;color:white;border:none;padding:9px 13px;border-radius:7px;font-weight:700;cursor:pointer;">📋 Copy label image</button>
            <a download="nutrition_facts_label.png" href="data:image/png;base64,{b64}" style="background:#f8fafc;color:#111827;text-decoration:none;border:1px solid #cbd5e1;padding:9px 13px;border-radius:7px;font-weight:700;">📥 Download PNG</a>
            <span id="status_{key}" style="font-size:12px;color:#475569;"></span>
        </div>
        <script>
        const btn_{key} = document.getElementById("copy_{key}");
        const status_{key} = document.getElementById("status_{key}");
        btn_{key}.onclick = async () => {{
            try {{
                if (!navigator.clipboard || !window.ClipboardItem) throw new Error("Image clipboard not supported");
                const binary = atob("{b64}");
                const bytes = new Uint8Array(binary.length);
                for (let i = 0; i < binary.length; i++) bytes[i] = binary.charCodeAt(i);
                const blob = new Blob([bytes], {{type: "image/png"}});
                await navigator.clipboard.write([new ClipboardItem({{"image/png": blob}})]);
                status_{key}.innerText = "✅ Copied image. Paste into Word, PowerPoint, Canva, etc.";
                btn_{key}.innerText = "✅ Copied";
                setTimeout(() => {{ btn_{key}.innerText = "📋 Copy label image"; }}, 1800);
            }} catch (err) {{
                status_{key}.innerText = "⚠️ Browser blocked image copy. Use Download PNG.";
                btn_{key}.innerText = "Copy blocked";
                setTimeout(() => {{ btn_{key}.innerText = "📋 Copy label image"; }}, 2200);
            }}
        }};
        </script>
    ''', height=76)


def create_zpl_from_panel(panel_text, label_size="2 x 4 in", dpi=203):
    size_map = {"2 x 4 in": (2, 4), "3 x 5 in": (3, 5), "4 x 6 in": (4, 6), "5 x 7 in": (5, 7)}
    w_in, h_in = size_map.get(label_size, (2, 4))
    w, h = int(w_in * dpi), int(h_in * dpi)
    lines = [re.sub(r"[^A-Za-z0-9 %/*.,:;()<>+-]", "", line.strip()) for line in (panel_text or "").splitlines() if line.strip()]
    z = ["^XA", f"^PW{w}", f"^LL{h}", "^CI28"]
    margin = max(int(0.08 * dpi), 16)
    y = margin
    z.append(f"^FO{margin},{margin}^GB{w-2*margin},{h-2*margin},3^FS")
    for i, line in enumerate(lines):
        if y > h - margin - 30:
            break
        if i == 0:
            z.append(f"^FO{margin+10},{y}^A0N,42,38^FD{line}^FS")
            y += 52
            z.append(f"^FO{margin+10},{y}^GB{w-2*margin-20},8,8^FS")
            y += 18
        elif line.lower().startswith("calories"):
            z.append(f"^FO{margin+10},{y}^A0N,38,34^FD{line}^FS")
            y += 46
            z.append(f"^FO{margin+10},{y}^GB{w-2*margin-20},8,8^FS")
            y += 16
        else:
            font_h = 24 if not line.startswith("*") else 17
            font_w = 21 if not line.startswith("*") else 15
            z.append(f"^FO{margin+12},{y}^A0N,{font_h},{font_w}^FD{line[:54]}^FS")
            y += font_h + 5
    z.append("^XZ")
    return "\n".join(z)


def create_batch_zpl_zip(saved_recipes, label_size="2 x 4 in", dpi=203):
    output = BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        for recipe in saved_recipes:
            name = recipe.get("name", "recipe")
            panel_text = recipe.get("nutrition_facts_panel") or nutrition_facts_panel_text(name, recipe.get("nutrition_per_serving", {}), recipe.get("servings", 1), recipe.get("serving_weight_g", 0), recipe.get("serving_size_label"))
            safe_name = re.sub(r"[^A-Za-z0-9_\-]+", "_", name).strip("_") or "recipe"
            zf.writestr(f"{safe_name}_nutrition_facts.zpl", create_zpl_from_panel(panel_text, label_size, dpi))
    return output.getvalue()


def create_batch_label_zip(saved_recipes, label_size="2 x 4 in", dpi=300):
    output = BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        for recipe in saved_recipes:
            name = recipe.get("name", "recipe")
            panel_text = recipe.get("nutrition_facts_panel") or nutrition_facts_panel_text(
                name,
                recipe.get("nutrition_per_serving", {}),
                recipe.get("servings", 1),
                recipe.get("serving_weight_g", 0),
                recipe.get("serving_size_label"),
            )
            png = create_nutrition_facts_png(panel_text, label_size, dpi)
            if png:
                safe_name = re.sub(r"[^A-Za-z0-9_\-]+", "_", name).strip("_") or "recipe"
                zf.writestr(f"{safe_name}_nutrition_facts_{label_size.replace(' ', '')}_{dpi}dpi.png", png)
    return output.getvalue()



tabs = st.tabs(["Dashboard", "Add Product", "Batch Upload Products", "Batch Upload Recipes", "Recipe Builder", "Saved Recipes"])

with tabs[0]:
    st.header("Dashboard")
    c1, c2, c3 = st.columns(3)
    c1.metric("Products", len(st.session_state.products))
    c2.metric("Recipe Items", len(st.session_state.recipe_items))
    c3.metric("Saved Recipes", len(st.session_state.saved_recipes))
    st.subheader("Customer + Sample Database")
    st.dataframe(product_table(st.session_state.products), use_container_width=True)
    st.info("USDA results require a Streamlit secret named USDA_API_KEY. Open Food Facts does not require a key.")

with tabs[1]:
    st.header("Add Customer Product / Ingredient")
    name = st.text_input("Product name")
    ingredients = st.text_area("Ingredients")
    allergens = st.text_input("Allergens", placeholder="milk, cereals containing gluten")
    c1, c2, c3, c4, c5 = st.columns(5)
    calories = c1.number_input("Calories", min_value=0.0, value=0.0)
    protein = c2.number_input("Protein g", min_value=0.0, value=0.0)
    fat = c3.number_input("Fat g", min_value=0.0, value=0.0)
    carbs = c4.number_input("Carbs g", min_value=0.0, value=0.0)
    salt = c5.number_input("Salt g", min_value=0.0, value=0.0)
    serving_note = st.text_input("Serving note", placeholder="Example: 150 cal per 100 g")
    if st.button("Save Product"):
        if name:
            st.session_state.products.append(normalize_product({
                "name": name,
                "source": "Customer",
                "calories": calories,
                "protein": protein,
                "fat": fat,
                "carbs": carbs,
                "salt": salt,
                "allergens": allergens or detect_allergens(ingredients),
                "ingredients": ingredients,
                "serving_note": serving_note or f"{calories:g} cal per serving",
            }))
            st.success("Saved")
        else:
            st.warning("Add a name first")

with tabs[2]:
    st.header("Batch Upload Product Specifications")
    st.caption("Upload multiple .txt/.pdf files or a .zip folder containing individual specification files. Each file becomes one searchable product/ingredient.")

    uploads = st.file_uploader(
        "Upload product specification files",
        type=["txt", "pdf", "zip"],
        accept_multiple_files=True,
    )

    if uploads:
        if st.button("Extract Products from Uploaded Files"):
            parsed, errors = parse_batch_uploads(uploads)
            st.session_state["batch_preview"] = parsed
            st.session_state["batch_errors"] = errors

    batch_preview = st.session_state.get("batch_preview", [])
    batch_errors = st.session_state.get("batch_errors", [])

    if batch_errors:
        st.warning("Some files need review:")
        for err in batch_errors:
            st.write("-", err)

    if batch_preview:
        st.success(f"Extracted {len(batch_preview)} product(s). Review below, then add them to the database.")
        st.dataframe(product_table(batch_preview), use_container_width=True)

        for i, product in enumerate(batch_preview):
            with st.expander(f"Review: {product['name']}"):
                st.write("Source file:", product.get("filename", ""))
                product["name"] = st.text_input("Product name", value=product["name"], key=f"batch_name_{i}")
                product["ingredients"] = st.text_area("Ingredients", value=product.get("ingredients", ""), key=f"batch_ing_{i}", height=90)
                product["allergens"] = st.text_input("Allergens", value=product.get("allergens", ""), key=f"batch_allergens_{i}")
                c1, c2, c3, c4, c5 = st.columns(5)
                product["calories"] = c1.number_input("Calories", value=float(product.get("calories", 0)), key=f"batch_cal_{i}")
                product["protein"] = c2.number_input("Protein g", value=float(product.get("protein", 0)), key=f"batch_pro_{i}")
                product["fat"] = c3.number_input("Fat g", value=float(product.get("fat", 0)), key=f"batch_fat_{i}")
                product["carbs"] = c4.number_input("Carbs g", value=float(product.get("carbs", 0)), key=f"batch_carb_{i}")
                product["salt"] = c5.number_input("Salt g", value=float(product.get("salt", 0)), key=f"batch_salt_{i}")

        if st.button("Add All Extracted Products to Database"):
            for product in batch_preview:
                product["source"] = "Customer"
                st.session_state.products.append(normalize_product(product))
            st.success(f"Added {len(batch_preview)} products to database search.")
            st.session_state["batch_preview"] = []
            st.session_state["batch_errors"] = []
            st.rerun()
    else:
        st.info("Tip: ZIP upload is the easiest way to upload a folder of product specs from your computer.")

with tabs[3]:
    st.header("Batch Upload Recipes")
    st.caption("Upload multiple recipe files or a .zip folder. Supported files: .txt, .pdf, .docx, .xlsx, .xls, .csv, and .zip.")

    recipe_uploads = st.file_uploader(
        "Upload recipe files",
        type=["txt", "pdf", "docx", "xlsx", "xls", "csv", "zip"],
        accept_multiple_files=True,
        key="recipe_batch_uploads",
    )

    if recipe_uploads:
        if st.button("Extract Recipes from Uploaded Files"):
            parsed_recipes, recipe_errors = parse_recipe_batch_uploads(recipe_uploads)
            st.session_state["recipe_batch_preview"] = parsed_recipes
            st.session_state["recipe_batch_errors"] = recipe_errors

    recipe_batch_preview = st.session_state.get("recipe_batch_preview", [])
    recipe_batch_errors = st.session_state.get("recipe_batch_errors", [])

    if recipe_batch_errors:
        st.warning("Some recipe files need review:")
        for err in recipe_batch_errors:
            st.write("-", err)

    if recipe_batch_preview:
        st.success(f"Extracted {len(recipe_batch_preview)} recipe(s). Review below, then add them to Saved Recipes.")
        for i, recipe in enumerate(recipe_batch_preview):
            with st.expander(f"Review Recipe: {recipe['name']}", expanded=(i == 0)):
                st.write("Source file:", recipe.get("filename", ""))
                recipe["name"] = st.text_input("Recipe name", value=recipe["name"], key=f"recipe_batch_name_{i}")
                recipe["servings"] = st.number_input("Servings", min_value=1, value=int(recipe.get("servings", 1)), key=f"recipe_batch_servings_{i}")
                recipe["uploaded_ingredients"] = st.text_area("Ingredient statement", value=recipe.get("uploaded_ingredients", ""), height=110, key=f"recipe_batch_ingredients_{i}")
                recipe["allergens"] = st.text_input("Allergens", value=recipe.get("allergens", ""), key=f"recipe_batch_allergens_{i}")
                c1, c2, c3, c4, c5 = st.columns(5)
                recipe["nutrition_per_serving"]["calories"] = c1.number_input("Calories / serving", value=float(recipe["nutrition_per_serving"].get("calories", 0)), key=f"recipe_batch_cal_{i}")
                recipe["nutrition_per_serving"]["protein"] = c2.number_input("Protein g / serving", value=float(recipe["nutrition_per_serving"].get("protein", 0)), key=f"recipe_batch_pro_{i}")
                recipe["nutrition_per_serving"]["fat"] = c3.number_input("Fat g / serving", value=float(recipe["nutrition_per_serving"].get("fat", 0)), key=f"recipe_batch_fat_{i}")
                recipe["nutrition_per_serving"]["carbs"] = c4.number_input("Carbs g / serving", value=float(recipe["nutrition_per_serving"].get("carbs", 0)), key=f"recipe_batch_carbs_{i}")
                recipe["nutrition_per_serving"]["salt"] = c5.number_input("Salt g / serving", value=float(recipe["nutrition_per_serving"].get("salt", 0)), key=f"recipe_batch_salt_{i}")
                allergen_text = "Contains: " + (recipe.get("allergens") or "No declarable allergens detected")
                recipe["label"] = f"""{recipe['name']}\n\nIngredients: {recipe.get('uploaded_ingredients', '')}\n\n{allergen_text}\n\nNutrition per serving:\nEnergy: {recipe['nutrition_per_serving']['calories']} kcal\nFat: {recipe['nutrition_per_serving']['fat']} g\nCarbohydrate: {recipe['nutrition_per_serving']['carbs']} g\nProtein: {recipe['nutrition_per_serving']['protein']} g\nSalt: {recipe['nutrition_per_serving']['salt']} g\n"""
                st.text_area("Generated label preview", value=recipe["label"], height=220, key=f"recipe_batch_label_{i}")

        if st.button("Add All Extracted Recipes to Saved Recipes"):
            for recipe in recipe_batch_preview:
                st.session_state.saved_recipes.append({
                    "name": recipe["name"],
                    "servings": recipe["servings"],
                    "items": recipe.get("items", []),
                    "label": recipe["label"],
                    "nutrition_per_serving": recipe["nutrition_per_serving"],
                    "uploaded_ingredients": recipe.get("uploaded_ingredients", ""),
                    "source_file": recipe.get("filename", ""),
                })
            st.success(f"Added {len(recipe_batch_preview)} recipes to Saved Recipes.")
            st.session_state["recipe_batch_preview"] = []
            st.session_state["recipe_batch_errors"] = []
            st.rerun()
    else:
        st.info("Tip: ZIP upload is the easiest way to upload a folder of recipe docs from your computer.")

with tabs[4]:
    st.header("Recipe Builder")
    st.subheader("Database Search")
    st.caption("Start typing to preview predicted matches from Customer/Sample, USDA, and Open Food Facts. No search button needed.")
    q = st.text_input("Search ingredients/products", placeholder="Type chicken, bun, cheese, flour...", label_visibility="collapsed")

    if q.strip() and len(q.strip()) < 3:
        st.caption("Type at least 3 characters to search external databases. Customer/sample predictions may appear sooner.")

    with st.spinner("Searching databases...") if len(q.strip()) >= 3 else st.spinner("Loading suggestions..."):
        results, source_counts = combined_database_search(q)

    st.caption(f"Results by source: Customer/Sample {source_counts['Customer/Sample']} | USDA {source_counts['USDA']} | Open Food Facts {source_counts['Open Food Facts']}")

    if q and not results:
        st.info("No matching products found. Add it in the Add Product tab.")
    else:
        if q:
            st.caption(f"Showing predicted matches for: {q}")
        else:
            st.caption("Suggested database items")

        st.write(f"Total results available: {len(results)}")
        if results:
            display_options = []
            for idx, product in enumerate(results):
                product = normalize_product(product)
                display_options.append(
                    f"{idx + 1}. [{product.get('source', 'Unknown')}] {product.get('name', 'Unnamed')} — {product.get('calories', 0):g} cal"
                )

            selected_label = st.selectbox(
                "Scrollable result list — choose an item",
                display_options,
                index=0,
                key=f"recipe_search_select_{q}",
            )
            selected_index = display_options.index(selected_label)
            selected_product = normalize_product(results[selected_index])

            action_cols = st.columns([2, 2, 6])
            with action_cols[0]:
                if st.button("+ Add Ingredient", key=f"dropdown_add_{selected_index}_{selected_product['source']}_{selected_product['name']}", use_container_width=True):
                    item = dict(selected_product)
                    item["amount"] = 1.0
                    item["unit"] = "serving"
                    item["waste_pct"] = 0.0
                    item["grams"] = item_grams(item)
                    st.session_state.recipe_items.append(item)
                    st.success(f"Added {selected_product['name']}")
                    st.rerun()
            with action_cols[1]:
                preview = st.toggle("Preview", key=f"dropdown_preview_{selected_index}_{selected_product['source']}_{selected_product['name']}")
            with action_cols[2]:
                st.caption(f"Selected: {selected_product['name']} • {selected_product.get('source', 'Unknown')} • {selected_product.get('calories', 0):g} cal")

            if preview:
                with st.container(border=True):
                    st.markdown(
                        f"**{selected_product['name']}** &nbsp; {badge_html(selected_product.get('source', 'Customer'))}",
                        unsafe_allow_html=True,
                    )
                    st.caption(selected_product.get("serving_note") or f"{selected_product.get('calories', 0):g} cal per serving")
                    if selected_product.get("allergens"):
                        st.warning(f"Allergens: {selected_product.get('allergens')}")
                    st.write("**Ingredients**")
                    st.write(selected_product.get("ingredients", "") or "Not available")
                    st.write("**Nutrition**")
                    st.json({
                        "calories": selected_product.get("calories", 0),
                        "protein_g": selected_product.get("protein", 0),
                        "fat_g": selected_product.get("fat", 0),
                        "carbs_g": selected_product.get("carbs", 0),
                        "salt_g": selected_product.get("salt", 0),
                    })
    st.divider()
    st.subheader("Current Recipe Items")
    if not st.session_state.recipe_items:
        st.info("Use Database Search above to add items to your recipe.")
    else:
        st.caption("Edit quantities, choose units, capture waste %, review calculated grams, or mark items for deletion.")

        editor_rows = []
        for idx, item in enumerate(st.session_state.recipe_items):
            item.setdefault("amount", item.get("qty", 1.0))
            item.setdefault("unit", "serving")
            item.setdefault("waste_pct", 0.0)
            item["grams"] = item_grams(item)
            editor_rows.append({
                "Delete": False,
                "Ingredient": item.get("name", ""),
                "Source": item.get("source", ""),
                "Quantity": safe_float(item.get("amount", 1.0)),
                "Unit": item.get("unit", "serving"),
                "Waste %": safe_float(item.get("waste_pct", 0.0)),
                "Grams": item.get("grams", 0.0),
                "Calories": safe_float(item.get("calories", 0.0)),
                "Allergens": item.get("allergens", ""),
            })

        edited_df = st.data_editor(
            pd.DataFrame(editor_rows),
            use_container_width=True,
            hide_index=True,
            column_config={
                "Delete": st.column_config.CheckboxColumn("Delete"),
                "Ingredient": st.column_config.TextColumn("Ingredient", disabled=True),
                "Source": st.column_config.TextColumn("Source", disabled=True),
                "Quantity": st.column_config.NumberColumn("Quantity", min_value=0.0, step=0.25),
                "Unit": st.column_config.SelectboxColumn("Unit", options=COMMON_UNITS),
                "Waste %": st.column_config.NumberColumn("Waste %", min_value=0.0, max_value=100.0, step=1.0),
                "Grams": st.column_config.NumberColumn("Grams", disabled=True, help="Approximate usable grams after waste. Count units need manual review."),
                "Calories": st.column_config.NumberColumn("Calories", disabled=True),
                "Allergens": st.column_config.TextColumn("Allergens", disabled=True),
            },
            key="recipe_items_editor",
        )

        updated_items = []
        for idx, row in edited_df.iterrows():
            if bool(row.get("Delete", False)):
                continue
            item = dict(st.session_state.recipe_items[idx])
            item["amount"] = safe_float(row.get("Quantity", 1.0))
            item["qty"] = item["amount"]
            item["unit"] = row.get("Unit", "serving")
            item["waste_pct"] = safe_float(row.get("Waste %", 0.0))
            item["grams"] = item_grams(item)
            updated_items.append(item)

        c_apply, c_clear = st.columns([1, 1])
        if c_apply.button("Apply recipe table changes"):
            st.session_state.recipe_items = updated_items
            st.success("Recipe items updated.")
            st.rerun()
        if c_clear.button("Clear recipe items"):
            st.session_state.recipe_items = []
            st.rerun()

        if any(item_grams(item) == 0 and item.get("unit") in ["each", "piece", "slice", "serving", "portion"] for item in updated_items):
            st.warning("Some count-based units cannot be converted to grams without item-specific weights. Review predominance and nutrition scaling before using labels commercially.")

        recipe_name = st.text_input("Recipe name")
        servings = st.number_input("Servings", min_value=1, value=1)

        st.markdown("**Serving options**")
        serving_option = st.selectbox(
            "Nutrition label basis",
            [
                "Per serving (FDA default)",
                "Per 100 g",
                "Per full recipe / container",
                "Custom serving weight (g)",
            ],
            help="FDA Nutrition Facts labels are normally per serving. Other options help with internal review or non-retail/export workflows.",
        )
        ss_col1, ss_col2, ss_col3 = st.columns([1, 1, 1])
        serving_size_value = ss_col1.number_input("Serving size value", min_value=0.0, value=1.0, step=0.25)
        serving_size_unit = ss_col2.selectbox("Serving size unit", COMMON_UNITS, index=COMMON_UNITS.index("serving") if "serving" in COMMON_UNITS else 0)
        custom_serving_weight_g = ss_col3.number_input("Custom serving weight (g)", min_value=0.0, value=0.0, step=1.0, disabled=serving_option != "Custom serving weight (g)")

        total, allergens, ingredient_list = totals(updated_items)
        total_weight_g = round(sum(item_grams(item) for item in updated_items), 2)
        serving_weight_g = round(total_weight_g / max(servings, 1), 2)
        per, serving_size_label, label_servings = calculate_label_nutrition(
            total,
            servings,
            total_weight_g,
            serving_option,
            serving_size_value,
            serving_size_unit,
            custom_serving_weight_g,
        )

        st.subheader("Nutrition Facts Panel")
        st.caption("Locked FDA-style label preview. Values shown use the selected serving option above.")
        panel_text = nutrition_facts_panel_text(recipe_name or "Recipe", per, label_servings, serving_weight_g, serving_size_label)

        label_size = st.session_state.get("current_label_size", DEFAULT_LABEL_SIZE)
        label_dpi = st.session_state.get("current_label_dpi", DEFAULT_LABEL_DPI)
        label_png = create_nutrition_facts_png(panel_text, label_size, label_dpi)

        if label_png:
            left_label_col, right_label_col = st.columns([1.15, 1])
            with left_label_col:
                st.image(label_png, caption="Final Nutrition Facts panel")
            with right_label_col:
                st.markdown("**Label actions**")
                image_clipboard_tools(label_png, key="copy_current_label_image")
                panel_pdf = create_nutrition_facts_pdf(recipe_name or "Recipe", panel_text)
                if panel_pdf:
                    st.download_button(
                        "Download Label PDF",
                        panel_pdf,
                        file_name=f"{(recipe_name or 'recipe').replace(' ', '_')}_nutrition_facts_panel.pdf",
                        mime="application/pdf",
                    )
                with st.expander("Advanced print/export settings", expanded=False):
                    size_options = ["2 x 4 in", "3 x 5 in", "4 x 6 in", "5 x 7 in"]
                    dpi_options = [203, 300, 600]
                    label_cols = st.columns([1, 1])
                    label_cols[0].selectbox(
                        "Label size",
                        size_options,
                        index=size_options.index(label_size) if label_size in size_options else 0,
                        key="current_label_size",
                    )
                    label_cols[1].selectbox(
                        "Print DPI",
                        dpi_options,
                        index=dpi_options.index(label_dpi) if label_dpi in dpi_options else 1,
                        key="current_label_dpi",
                    )
                    zpl_text = create_zpl_from_panel(panel_text, st.session_state.get("current_label_size", DEFAULT_LABEL_SIZE), 203)
                    st.download_button(
                        "Download Zebra/ZPL label file",
                        zpl_text,
                        file_name=f"{(recipe_name or 'recipe').replace(' ', '_')}_nutrition_facts.zpl",
                        mime="text/plain",
                    )
                    st.caption("Changing size/DPI refreshes the panel after Streamlit reruns.")
                with st.expander("Show raw text version", expanded=False):
                    st.text_area("Nutrition Facts text", panel_text, height=320)
        else:
            st.error("Nutrition Facts image could not be generated. Confirm Pillow is listed in requirements.txt.")
            st.text_area("Fallback Nutrition Facts text", panel_text, height=320)

        st.subheader("Ingredient statement")
        st.text_area("Ingredient list", ingredient_list, height=100)
        st.subheader("Allergen declaration")
        allergen_text = "Contains: " + (", ".join(allergens) if allergens else "No declarable allergens detected")
        st.text_area("Allergens", allergen_text, height=80)

        label = f"""{recipe_name or 'Recipe'}\n\nIngredients: {ingredient_list}\n\n{allergen_text}\n\n{panel_text}"""
        if st.button("Save Recipe"):
            if recipe_name:
                st.session_state.recipe_items = updated_items
                st.session_state.saved_recipes.append({"name": recipe_name, "servings": label_servings, "original_servings": servings, "serving_option": serving_option, "serving_size_value": serving_size_value, "serving_size_unit": serving_size_unit, "serving_size_label": serving_size_label, "items": list(updated_items), "label": label, "nutrition_per_serving": per, "nutrition_facts_panel": panel_text, "serving_weight_g": serving_weight_g})
                st.success("Recipe saved")
            else:
                st.warning("Add a recipe name")

with tabs[5]:
    st.header("Saved Recipes + Full Nutrition Exports")

    if not st.session_state.saved_recipes:
        st.info("No saved recipes yet")
    else:
        st.subheader("Saved Recipe Labels")
        for r in st.session_state.saved_recipes:
            with st.expander(r["name"]):
                st.text_area("Label", r["label"], height=220)
                st.download_button("Download label text", r["label"], file_name=f"{r['name']}_label.txt")

                st.subheader("Nutrition Facts Panel")
                panel_per = r.get("nutrition_per_serving", {})
                panel_text = r.get("nutrition_facts_panel") or nutrition_facts_panel_text(r.get("name", "Recipe"), panel_per, r.get("servings", 1), r.get("serving_weight_g", 0))
                safe_recipe_key = re.sub(r'[^A-Za-z0-9_]+', '_', r['name'])
                with st.expander("Advanced print/export settings", expanded=False):
                    saved_cols = st.columns([1, 1])
                    saved_cols[0].selectbox("Label size", ["2 x 4 in", "3 x 5 in", "4 x 6 in", "5 x 7 in"], index=0, key=f"saved_label_size_{safe_recipe_key}")
                    saved_cols[1].selectbox("Print DPI", [203, 300, 600], index=1, key=f"saved_label_dpi_{safe_recipe_key}")
                saved_size = st.session_state.get(f"saved_label_size_{safe_recipe_key}", DEFAULT_LABEL_SIZE)
                saved_dpi = st.session_state.get(f"saved_label_dpi_{safe_recipe_key}", DEFAULT_LABEL_DPI)
                saved_png = create_nutrition_facts_png(panel_text, saved_size, saved_dpi)
                if saved_png:
                    st.image(saved_png, caption="Nutrition Facts label preview")
                    image_clipboard_tools(saved_png, key=f"copy_saved_label_image_{safe_recipe_key}")
                    with st.expander("PDF / Zebra printer export", expanded=False):
                        panel_pdf = create_nutrition_facts_pdf(r.get("name", "Recipe"), panel_text)
                        if panel_pdf:
                            st.download_button("Download Nutrition Facts PDF", panel_pdf, file_name=f"{r['name']}_nutrition_facts.pdf", mime="application/pdf")
                        saved_zpl = create_zpl_from_panel(panel_text, saved_size, 203)
                        st.download_button("Download Zebra/ZPL label", saved_zpl, file_name=f"{r['name']}_nutrition_facts.zpl", mime="text/plain")

        st.divider()
        st.subheader("Batch Print-Ready Label Export")
        with st.expander("Batch print/export settings", expanded=False):
            batch_cols = st.columns([1, 1])
            batch_cols[0].selectbox("Batch label size", ["2 x 4 in", "3 x 5 in", "4 x 6 in", "5 x 7 in"], index=0, key="batch_label_size")
            batch_cols[1].selectbox("Batch print DPI", [203, 300, 600], index=1, key="batch_label_dpi")
        batch_size = st.session_state.get("batch_label_size", DEFAULT_LABEL_SIZE)
        batch_dpi = st.session_state.get("batch_label_dpi", DEFAULT_LABEL_DPI)
        batch_zip = create_batch_label_zip(st.session_state.saved_recipes, batch_size, batch_dpi)
        st.download_button("Download all saved labels as PNG ZIP", batch_zip, file_name="nutrition_facts_labels_png.zip", mime="application/zip")
        batch_zpl_zip = create_batch_zpl_zip(st.session_state.saved_recipes, batch_size, 203)
        st.download_button("Download all saved labels as Zebra/ZPL ZIP", batch_zpl_zip, file_name="nutrition_facts_labels_zpl.zip", mime="application/zip")

        st.divider()
        st.subheader("Full Nutrition Export")
        st.caption("Exports follow the attached workbook structure: Nutritionals, Allergens, Ingredients, and GF/Vegan/Halal review sheets.")

        export_df = build_nutrition_export_dataframe(st.session_state.saved_recipes)
        st.dataframe(export_df, use_container_width=True)

        excel_data = create_excel_export(st.session_state.saved_recipes)
        st.download_button(
            "Download Excel nutrition workbook",
            data=excel_data,
            file_name="full_nutrition_export.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        pdf_data = create_pdf_export(st.session_state.saved_recipes)
        if pdf_data:
            st.download_button(
                "Download PDF nutrition report",
                data=pdf_data,
                file_name="full_nutrition_export.pdf",
                mime="application/pdf",
            )
        else:
            st.warning("PDF export requires reportlab. Make sure reportlab is included in requirements.txt.")

